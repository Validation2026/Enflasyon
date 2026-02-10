# GEREKLİ KÜTÜPHANELER:
# pip install streamlit-lottie python-docx plotly pandas xlsxwriter matplotlib requests

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from bs4 import BeautifulSoup
import re
from datetime import datetime, timedelta
import time
import json
from github import Github
from io import BytesIO
import zipfile
import base64
import requests
import streamlit.components.v1 as components
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit_lottie import st_lottie

# --- 1. AYARLAR VE TEMA YÖNETİMİ ---
st.set_page_config(
    page_title="Piyasa Monitörü | Pro Terminal",
    layout="wide",
    page_icon="💠",
    initial_sidebar_state="expanded"
)

# --- CSS MOTORU (PREMIUM FINTECH TEMA) ---
def apply_theme():
    if 'plotly_template' not in st.session_state:
        st.session_state.plotly_template = "plotly_dark"

    final_css = """
    <style>
        /* Fontları Yükle */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
        @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@500;700&display=swap');

        :root {
            /* Renk Paleti - Premium Dark */
            --bg-app: #0e1117; /* Ana arka plan */
            --bg-surface-1: #161b22; /* Yan menü ve kart zeminleri */
            --bg-surface-2: #21262d; /* Kart içi vurgular, inputlar */
            --border-subtle: rgba(240, 246, 252, 0.1);
            --text-primary: #f0f6fc;
            --text-secondary: #8b949e;
            --accent-blue: #2563eb; /* Daha kurumsal bir mavi */
            --accent-green: #2ea043;
            --accent-red: #da3633;
        }

        /* --- Temel Yapı --- */
        .stApp {
            background-color: var(--bg-app);
            font-family: 'Inter', sans-serif;
            color: var(--text-primary);
        }

        /* Header'ı gizle (daha temiz görünüm için) */
        header {visibility: hidden;}
        [data-testid="stHeader"] { visibility: hidden; height: 0px; }

        /* Sidebar Styling */
        section[data-testid="stSidebar"] {
            background-color: var(--bg-surface-1);
            border-right: 1px solid var(--border-subtle);
        }
        
        /* --- GELİŞMİŞ NAVİGASYON (Segmented Control) --- */
        [data-testid="stRadio"] {
            background-color: var(--bg-surface-1);
            padding: 4px;
            border-radius: 12px;
            border: 1px solid var(--border-subtle);
            display: inline-block; /* İçeriği kadar yer kaplasın */
        }

        [data-testid="stRadio"] > div[role="radiogroup"] {
            display: flex;
            gap: 4px;
            background: transparent;
        }

        [data-testid="stRadio"] label {
            flex: 1;
            text-align: center;
            padding: 8px 16px !important;
            border-radius: 8px !important;
            font-family: 'Inter', sans-serif;
            font-weight: 600 !important;
            font-size: 13px !important;
            color: var(--text-secondary) !important;
            background: transparent !important;
            border: none !important;
            transition: all 0.2s ease;
            display: flex;
            align-items: center;
            justify-content: center;
        }

        /* Seçili Olmayan Hover Durumu */
        [data-testid="stRadio"] label:hover {
            color: var(--text-primary) !important;
            background-color: rgba(255,255,255,0.03) !important;
        }

        /* Seçili Durum */
        [data-testid="stRadio"] label[data-checked="true"] {
            background-color: var(--accent-blue) !important;
            color: white !important;
            box-shadow: 0 2px 8px rgba(37, 99, 235, 0.3);
        }
        
        /* Radyo düğmesinin yuvarlak kısmını gizle */
        [data-testid="stRadio"] div[role="radiogroup"] > :first-child {
             display: none;
        }

        /* --- KART TASARIMLARI (Clean Surface) --- */
        .kpi-card {
            background-color: var(--bg-surface-1);
            border: 1px solid var(--border-subtle);
            border-radius: 16px;
            padding: 24px;
            display: flex;
            flex-direction: column;
            transition: border-color 0.3s, box-shadow 0.3s;
            height: 100%;
        }
        
        .kpi-card:hover {
            border-color: var(--accent-blue);
            box-shadow: 0 8px 24px rgba(0,0,0,0.2);
        }

        .kpi-title {
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--text-secondary);
            font-weight: 600;
            margin-bottom: 12px;
        }

        .kpi-value {
            font-family: 'JetBrains Mono', monospace;
            font-size: 36px;
            font-weight: 700;
            color: var(--text-primary);
            margin-bottom: 8px;
            letter-spacing: -0.5px;
        }
        
        .kpi-sub {
            font-size: 12px;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 6px;
        }
        .trend-up { color: var(--accent-green); }
        .trend-down { color: var(--accent-red); }
        .trend-flat { color: var(--text-secondary); }

        /* --- TICKER BANDI (Daha sade) --- */
        .ticker-container {
            background: var(--bg-surface-1);
            border-top: 1px solid var(--border-subtle);
            border-bottom: 1px solid var(--border-subtle);
            padding: 12px 0;
            overflow: hidden;
            white-space: nowrap;
            margin-bottom: 24px;
        }
        .ticker-track {
            display: inline-block;
            animation: marquee 60s linear infinite;
        }
        .ticker-item {
            font-family: 'JetBrains Mono', monospace;
            font-size: 13px;
            margin: 0 15px;
        }
        @keyframes marquee { 0% { transform: translateX(0); } 100% { transform: translateX(-100%); } }

        /* --- FİYAT GRID KARTLARI (Grid View) --- */
        .grid-card {
            background: var(--bg-surface-1);
            border: 1px solid var(--border-subtle);
            border-radius: 12px;
            padding: 16px;
            transition: all 0.2s ease;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            height: 130px;
        }
        .grid-card:hover {
            transform: translateY(-3px);
            border-color: rgba(255,255,255,0.2);
            background: var(--bg-surface-2);
        }
        .grid-title {
            font-size: 13px;
            font-weight: 500;
            color: var(--text-primary);
            line-height: 1.3;
            overflow: hidden; display: -webkit-box; -webkit-line-clamp: 2; -webkit-box-orient: vertical;
        }
        .grid-price-row {
            display: flex;
            justify-content: space-between;
            align-items: flex-end;
            margin-top: auto;
        }
        .grid-price {
            font-family: 'JetBrains Mono', monospace;
            font-size: 18px;
            font-weight: 700;
        }
        .grid-badge {
            font-size: 12px;
            font-weight: 700;
            padding: 4px 10px;
            border-radius: 6px;
        }
        .badge-up { background: rgba(46, 160, 67, 0.15); color: var(--accent-green); }
        .badge-down { background: rgba(218, 54, 51, 0.15); color: var(--accent-red); }
        .badge-flat { background: rgba(139, 148, 158, 0.15); color: var(--text-secondary); }

        /* --- BİLEŞEN ÖZELLEŞTİRMELERİ --- */
        /* Butonlar */
        div.stButton > button {
            background-color: var(--accent-blue);
            color: white;
            border: none;
            padding: 10px 24px;
            font-weight: 600;
            border-radius: 8px;
            transition: background-color 0.2s;
        }
        div.stButton > button:hover {
            background-color: #1d4ed8; /* Biraz daha koyu mavi */
            box-shadow: 0 4px 12px rgba(37, 99, 235, 0.2);
        }
        div.stButton > button:active {
            transform: scale(0.98);
        }

        /* Inputlar ve Selectboxlar */
        div[data-baseweb="select"] > div,
        div[data-baseweb="input"] > div,
        div[data-baseweb="base-input"] {
            background-color: var(--bg-surface-2) !important;
            border-color: var(--border-subtle) !important;
            border-radius: 8px !important;
            color: var(--text-primary) !important;
        }

        /* Expander (Genişletici) */
        .streamlit-expanderHeader {
            background-color: var(--bg-surface-1);
            border-radius: 8px;
        }
        
        /* Progress Bar */
        .stProgress > div > div > div {
            background-color: var(--accent-blue);
        }
    </style>
    """
    st.markdown(final_css, unsafe_allow_html=True)

apply_theme()

# --- 2. GITHUB & VERİ MOTORU (DEĞİŞMEDİ) ---
EXCEL_DOSYASI = "TUFE_Konfigurasyon.xlsx"
FIYAT_DOSYASI = "Fiyat_Veritabani.xlsx"
SAYFA_ADI = "Madde_Sepeti"

def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None

# --- 3. RAPOR MOTORU (DEĞİŞMEDİ) ---
def create_word_report(text_content, tarih, df_analiz=None):
    try:
        doc = Document()
        matplotlib.use('Agg')
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        head = doc.add_heading(f'PİYASA GÖRÜNÜM RAPORU', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subhead = doc.add_paragraph(f'Rapor Tarihi: {tarih}')
        subhead.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph("")
        paragraphs = text_content.split('\n')
        for p_text in paragraphs:
            if not p_text.strip(): continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parts = p_text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: 
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 50, 100) 
        if df_analiz is not None and not df_analiz.empty:
            doc.add_page_break()
            doc.add_heading('EKLER: GÖRSEL ANALİZLER', 1)
            doc.add_paragraph("")
            try:
                if 'Fark' in df_analiz.columns:
                    data = pd.to_numeric(df_analiz['Fark'], errors='coerce').dropna() * 100
                    if not data.empty:
                        fig, ax = plt.subplots(figsize=(6, 4))
                        ax.hist(data, bins=20, color='#2563eb', edgecolor='#161b22', alpha=0.8)
                        ax.set_title(f"Fiyat Değişim Dağılımı (%) - {tarih}", fontsize=12, fontweight='bold')
                        ax.set_facecolor('#161b22')
                        fig.patch.set_facecolor('#161b22')
                        ax.tick_params(axis='x', colors='white')
                        ax.tick_params(axis='y', colors='white')
                        ax.title.set_color('white')
                        memfile = BytesIO()
                        plt.savefig(memfile, format='png', dpi=100, bbox_inches='tight')
                        plt.close(fig)
                        doc.add_picture(memfile, width=Inches(5.5))
                        memfile.close()
                        doc.add_paragraph("Grafik 1: Ürünlerin fiyat değişim oranlarına göre dağılımı.")
            except Exception:
                pass
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        return BytesIO()

# --- 4. GITHUB İŞLEMLERİ (DEĞİŞMEDİ) ---
@st.cache_resource
def get_github_connection():
    try:
        return Github(st.secrets["github"]["token"])
    except:
        return None

def get_github_repo():
    g = get_github_connection()
    if g:
        return g.get_repo(st.secrets["github"]["repo_name"])
    return None

@st.cache_data(ttl=600, show_spinner=False)
def github_excel_oku(dosya_adi, sayfa_adi=None):
    repo = get_github_repo()
    if not repo: return pd.DataFrame()
    try:
        c = repo.get_contents(dosya_adi, ref=st.secrets["github"]["branch"])
        if sayfa_adi:
            df = pd.read_excel(BytesIO(c.decoded_content), sheet_name=sayfa_adi, dtype=str)
        else:
            df = pd.read_excel(BytesIO(c.decoded_content), dtype=str)
        return df
    except:
        return pd.DataFrame()

def github_excel_guncelle(df_yeni, dosya_adi):
    repo = get_github_repo()
    if not repo: return "Repo Yok"
    try:
        try:
            c = repo.get_contents(dosya_adi, ref=st.secrets["github"]["branch"])
            old = pd.read_excel(BytesIO(c.decoded_content), dtype=str)
            yeni_tarih = str(df_yeni['Tarih'].iloc[0])
            old = old[~((old['Tarih'].astype(str) == yeni_tarih) & (old['Kod'].isin(df_yeni['Kod'])))]
            final = pd.concat([old, df_yeni], ignore_index=True)
        except:
            c = None; final = df_yeni
        out = BytesIO()
        with pd.ExcelWriter(out, engine='openpyxl') as w:
            final.to_excel(w, index=False, sheet_name='Fiyat_Log')
        msg = f"Data Update"
        if c:
            repo.update_file(c.path, msg, out.getvalue(), c.sha, branch=st.secrets["github"]["branch"])
        else:
            repo.create_file(dosya_adi, msg, out.getvalue(), branch=st.secrets["github"]["branch"])
        return "OK"
    except Exception as e:
        return str(e)

# --- 5. RESMİ ENFLASYON (CACHED) (DEĞİŞMEDİ) ---
@st.cache_data(ttl=3600, show_spinner=False)
def get_official_inflation():
    api_key = st.secrets.get("evds", {}).get("api_key")
    if not api_key: return None, "API Key Yok"
    start_date = (datetime.now() - timedelta(days=365)).strftime("%d-%m-%Y")
    end_date = datetime.now().strftime("%d-%m-%Y")
    url = f"https://evds2.tcmb.gov.tr/service/evds/series=TP.FG.J0&startDate={start_date}&endDate={end_date}&type=json"
    headers = {'User-Agent': 'Mozilla/5.0', 'key': api_key, 'Accept': 'application/json'}
    try:
        url_with_key = f"{url}&key={api_key}"
        res = requests.get(url_with_key, headers=headers, timeout=10, verify=False)
        if res.status_code == 200:
            data = res.json()
            if "items" in data:
                df_evds = pd.DataFrame(data["items"])
                df_evds = df_evds[['Tarih', 'TP_FG_J0']]
                df_evds.columns = ['Tarih', 'Resmi_TUFE']
                df_evds['Tarih'] = pd.to_datetime(df_evds['Tarih'] + "-01", format="%Y-%m-%d")
                df_evds['Resmi_TUFE'] = pd.to_numeric(df_evds['Resmi_TUFE'], errors='coerce')
                return df_evds, "OK"
        return None, "Hata"
    except Exception as e:
        return None, str(e)

# --- 6. SCRAPER YARDIMCILARI (DEĞİŞMEDİ) ---
def temizle_fiyat(t):
    if not t: return None
    t = str(t).replace('TL', '').replace('₺', '').strip()
    t = t.replace('.', '').replace(',', '.') if ',' in t and '.' in t else t.replace(',', '.')
    try:
        return float(re.sub(r'[^\d.]', '', t))
    except:
        return None

def kod_standartlastir(k): return str(k).replace('.0', '').strip().zfill(7)

def fiyat_bul_siteye_gore(soup, url):
    fiyat = 0; kaynak = ""; domain = url.lower() if url else ""
    # Basit Regex ve CSS arama
    if m := re.search(r'(\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)\s*(?:TL|₺)', soup.get_text()[:5000]):
        if v := temizle_fiyat(m.group(1)): fiyat = v; kaynak = "Regex"
    return fiyat, kaynak

def html_isleyici(progress_callback):
    repo = get_github_repo()
    if not repo: return "GitHub Bağlantı Hatası"
    progress_callback(0.05) 
    try:
        df_conf = github_excel_oku(EXCEL_DOSYASI, SAYFA_ADI)
        df_conf.columns = df_conf.columns.str.strip()
        kod_col = next((c for c in df_conf.columns if c.lower() == 'kod'), None)
        url_col = next((c for c in df_conf.columns if c.lower() == 'url'), None)
        ad_col = next((c for c in df_conf.columns if 'ad' in c.lower()), 'Madde adı')
        if not kod_col or not url_col: return "Hata: Excel sütunları eksik."
        df_conf['Kod'] = df_conf[kod_col].astype(str).apply(kod_standartlastir)
        url_map = {str(row[url_col]).strip(): row for _, row in df_conf.iterrows() if pd.notna(row[url_col])}
        veriler = []
        islenen_kodlar = set()
        bugun = datetime.now().strftime("%Y-%m-%d")
        simdi = datetime.now().strftime("%H:%M")
        
        progress_callback(0.10)
        contents = repo.get_contents("", ref=st.secrets["github"]["branch"])
        zip_files = [c for c in contents if c.name.endswith(".zip") and c.name.startswith("Bolum")]
        total_zips = len(zip_files)
        
        for i, zip_file in enumerate(zip_files):
            current_progress = 0.10 + (0.80 * ((i + 1) / max(1, total_zips)))
            progress_callback(current_progress)
            try:
                blob = repo.get_git_blob(zip_file.sha)
                zip_data = base64.b64decode(blob.content)
                with zipfile.ZipFile(BytesIO(zip_data)) as z:
                    for file_name in z.namelist():
                        if not file_name.endswith(('.html', '.htm')): continue
                        with z.open(file_name) as f:
                            raw = f.read().decode("utf-8", errors="ignore")
                            soup = BeautifulSoup(raw, 'html.parser')
                            found_url = None
                            if c := soup.find("link", rel="canonical"): found_url = c.get("href")
                            if found_url and str(found_url).strip() in url_map:
                                target = url_map[str(found_url).strip()]
                                if target['Kod'] in islenen_kodlar: continue
                                fiyat, kaynak = fiyat_bul_siteye_gore(soup, target[url_col])
                                if fiyat > 0:
                                    veriler.append({"Tarih": bugun, "Zaman": simdi, "Kod": target['Kod'],
                                                    "Madde_Adi": target[ad_col], "Fiyat": float(fiyat),
                                                    "Kaynak": kaynak, "URL": target[url_col]})
                                    islenen_kodlar.add(target['Kod'])
            except: pass
        
        progress_callback(0.95)
        if veriler:
            return github_excel_guncelle(pd.DataFrame(veriler), FIYAT_DOSYASI)
        else:
            return "Veri bulunamadı."
    except Exception as e:
        return f"Hata: {str(e)}"

# --- 7. STATİK ANALİZ MOTORU (DEĞİŞMEDİ) ---
def generate_detailed_static_report(df_analiz, tarih, enf_genel, enf_gida, gun_farki, tahmin, ad_col, agirlik_col):
    df_clean = df_analiz.dropna(subset=['Fark'])
    toplam_urun = len(df_clean)
    artanlar = df_clean[df_clean['Fark'] > 0]
    dusenler = df_clean[df_clean['Fark'] < 0]
    sabitler = df_clean[df_clean['Fark'] == 0]
    artan_sayisi = len(artanlar)
    yayilim_orani = (artan_sayisi / toplam_urun) * 100 if toplam_urun > 0 else 0
    inc = df_clean.sort_values('Fark', ascending=False).head(5)
    dec = df_clean.sort_values('Fark', ascending=True).head(5)
    inc_str = "\n".join([f"   🔴 %{row['Fark']*100:5.2f} | {row[ad_col]}" for _, row in inc.iterrows()])
    dec_str = "\n".join([f"   🟢 %{abs(row['Fark']*100):5.2f} | {row[ad_col]}" for _, row in dec.iterrows()])

    text = f"""
**PİYASA GÖRÜNÜM RAPORU**
**Tarih:** {tarih}

**1. 📊 ANA GÖSTERGELER**
-----------------------------------------
**GENEL ENFLASYON** : **%{enf_genel:.2f}**
**GIDA ENFLASYONU** : **%{enf_gida:.2f}**
**AY SONU TAHMİNİ** : **%{tahmin:.2f}**
-----------------------------------------

**2. 🔎 PİYASA RÖNTGENİ**
**Fiyat Hareketleri:**
   🔺 **Zamlanan Ürün:** {artan_sayisi} adet
   🔻 **İndirimli Ürün:** {len(dusenler)} adet
   ➖ **Fiyatı Değişmeyen:** {len(sabitler)} adet

**Sepet Yayılımı:**
   Her 100 üründen **{int(yayilim_orani)}** tanesinde fiyat artışı tespit edilmiştir.

**3. ⚡ DİKKAT ÇEKEN ÜRÜNLER**

**▲ Yüksek Artışlar (Cep Yakanlar)**
{inc_str}

**▼ Fiyat Düşüşleri (Fırsatlar)**
{dec_str}

**4. 💡 SONUÇ**
Tahmin modelimiz, ay sonu kapanışının **%{tahmin:.2f}** bandında olacağını öngörmektedir.

---
*Otomatik Rapor Sistemi | Validasyon Müdürlüğü*
"""
    return text.strip()

# --- YENİ YARDIMCI FONKSİYONLAR ---
def style_chart(fig, is_sunburst=False):
    # Yeni, profesyonel koyu tema
    layout_args = dict(
        template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Inter, sans-serif", color="#8b949e", size=12),
        margin=dict(l=10, r=10, t=40, b=10),
        hoverlabel=dict(
            bgcolor="#161b22", 
            bordercolor="rgba(240, 246, 252, 0.1)", 
            font=dict(family="JetBrains Mono", color="#f0f6fc", size=13)
        ),
        title_font=dict(size=16, color="#f0f6fc", family="Inter, sans-serif", weight=600),
    )
    
    if not is_sunburst:
        # Eksenleri daha belirgin ama inceltilmiş yap
        axis_style = dict(
            showgrid=True,
            gridcolor="rgba(240, 246, 252, 0.06)", # Çok hafif grid
            gridwidth=1,
            zeroline=True,
            zerolinecolor="rgba(240, 246, 252, 0.15)",
            showline=True,
            linecolor="rgba(240, 246, 252, 0.15)",
            tickfont=dict(color="#8b949e", size=11),
        )
        layout_args.update(dict(
            xaxis=axis_style,
            yaxis=axis_style
        ))
        
    fig.update_layout(**layout_args)
    
    # Çizgi grafiklerinde modernlik için yumuşatma (spline)
    if not is_sunburst:
        for data in fig.data:
            if data.type == 'scatter' and data.mode == 'lines+markers':
                data.line.shape = 'spline' 
                data.line.width = 3
                data.marker.size = 8
                data.marker.line.width = 2
                data.marker.line.color = '#0e1117' # Marker etrafına outline

    return fig

# --- 9. VERİ VE HESAPLAMA MOTORLARI (CACHE) (DEĞİŞMEDİ) ---

# 1. VERİ GETİR
@st.cache_data(ttl=600, show_spinner=False)
def verileri_getir_cache():
    df_f = github_excel_oku(FIYAT_DOSYASI)
    df_s = github_excel_oku(EXCEL_DOSYASI, SAYFA_ADI)
    if df_f.empty or df_s.empty: return None, None, None

    df_f['Tarih_DT'] = pd.to_datetime(df_f['Tarih'], errors='coerce')
    df_f = df_f.dropna(subset=['Tarih_DT']).sort_values('Tarih_DT')
    df_f['Tarih_Str'] = df_f['Tarih_DT'].dt.strftime('%Y-%m-%d')
    raw_dates = df_f['Tarih_Str'].unique().tolist()

    df_s.columns = df_s.columns.str.strip()
    kod_col = next((c for c in df_s.columns if c.lower() == 'kod'), 'Kod')
    ad_col = next((c for c in df_s.columns if 'ad' in c.lower()), 'Madde_Adi')
    df_f['Kod'] = df_f['Kod'].astype(str).apply(kod_standartlastir)
    df_s['Kod'] = df_s[kod_col].astype(str).apply(kod_standartlastir)
    df_s = df_s.drop_duplicates(subset=['Kod'], keep='first')
    df_f['Fiyat'] = pd.to_numeric(df_f['Fiyat'], errors='coerce')
    df_f = df_f[df_f['Fiyat'] > 0]
    
    pivot = df_f.pivot_table(index='Kod', columns='Tarih_Str', values='Fiyat', aggfunc='mean')
    pivot = pivot.ffill(axis=1).bfill(axis=1).reset_index()
    if pivot.empty: return None, None, None

    if 'Grup' not in df_s.columns:
        grup_map = {"01": "Gıda", "02": "Alkol-Tütün", "03": "Giyim", "04": "Konut"}
        df_s['Grup'] = df_s['Kod'].str[:2].map(grup_map).fillna("Diğer")

    df_analiz_base = pd.merge(df_s, pivot, on='Kod', how='left')
    return df_analiz_base, raw_dates, ad_col

# 2. HESAPLAMA YAP (CACHED)
@st.cache_data(show_spinner=False)
def hesapla_metrikler(df_analiz_base, secilen_tarih, gunler, tum_gunler_sirali, ad_col, agirlik_col, baz_col, aktif_agirlik_col, son):
    df_analiz = df_analiz_base.copy()
    for col in gunler: df_analiz[col] = pd.to_numeric(df_analiz[col], errors='coerce')
    dt_son = datetime.strptime(son, '%Y-%m-%d')
    if baz_col in df_analiz.columns: df_analiz[baz_col] = df_analiz[baz_col].fillna(df_analiz[son])
    df_analiz[aktif_agirlik_col] = pd.to_numeric(df_analiz.get(aktif_agirlik_col, 0), errors='coerce').fillna(0)
    gecerli_veri = df_analiz[df_analiz[aktif_agirlik_col] > 0].copy()
    
    def geo_mean(row):
        vals = [x for x in row if isinstance(x, (int, float)) and x > 0]
        return np.exp(np.mean(np.log(vals))) if vals else np.nan

    bu_ay_str = f"{dt_son.year}-{dt_son.month:02d}"
    bu_ay_cols = [c for c in gunler if c.startswith(bu_ay_str)]
    if not bu_ay_cols: bu_ay_cols = [son]
    
    gecerli_veri['Aylik_Ortalama'] = gecerli_veri[bu_ay_cols].apply(geo_mean, axis=1)
    gecerli_veri = gecerli_veri.dropna(subset=['Aylik_Ortalama', baz_col])

    enf_genel = 0.0; enf_gida = 0.0
    if not gecerli_veri.empty:
        w = gecerli_veri[aktif_agirlik_col]
        p_rel = gecerli_veri['Aylik_Ortalama'] / gecerli_veri[baz_col]
        if w.sum() > 0: enf_genel = (w * p_rel).sum() / w.sum() * 100 - 100
        
        gida_df = gecerli_veri[gecerli_veri['Kod'].astype(str).str.startswith("01")]
        if not gida_df.empty and gida_df[aktif_agirlik_col].sum() > 0:
            enf_gida = ((gida_df[aktif_agirlik_col] * (gida_df['Aylik_Ortalama']/gida_df[baz_col])).sum() / gida_df[aktif_agirlik_col].sum() * 100) - 100
            
        df_analiz['Fark'] = 0.0
        df_analiz.loc[gecerli_veri.index, 'Fark'] = (gecerli_veri['Aylik_Ortalama'] / gecerli_veri[baz_col]) - 1
        df_analiz['Fark_Yuzde'] = df_analiz['Fark'] * 100
    
    gun_farki = 0
    if len(gunler) >= 2:
        onceki_gun = gunler[-2]
        df_analiz['Gunluk_Degisim'] = (df_analiz[son] / df_analiz[onceki_gun].replace(0, np.nan)) - 1
    else:
        df_analiz['Gunluk_Degisim'] = 0
        onceki_gun = son

    month_end_forecast = 0.0
    target_fixed = f"{dt_son.year}-{dt_son.month:02d}-31"
    fixed_cols = [c for c in tum_gunler_sirali if c.startswith(bu_ay_str) and c <= target_fixed]
    if fixed_cols and not gecerli_veri.empty:
        gecerli_veri['Fixed_Ort'] = gecerli_veri[fixed_cols].apply(geo_mean, axis=1)
        gecerli_t = gecerli_veri.dropna(subset=['Fixed_Ort'])
        if not gecerli_t.empty and gecerli_t[aktif_agirlik_col].sum() > 0:
             month_end_forecast = ((gecerli_t[aktif_agirlik_col] * (gecerli_t['Fixed_Ort']/gecerli_t[baz_col])).sum() / gecerli_t[aktif_agirlik_col].sum() * 100) - 100

    resmi_aylik_degisim = 0.0
    try:
        df_resmi, _ = get_official_inflation()
        if df_resmi is not None and not df_resmi.empty:
             df_resmi = df_resmi.sort_values('Tarih')
             if len(df_resmi) >= 2:
                 son_endeks = df_resmi.iloc[-1]['Resmi_TUFE']
                 onceki_endeks = df_resmi.iloc[-2]['Resmi_TUFE']
                 resmi_aylik_degisim = ((son_endeks / onceki_endeks) - 1) * 100
    except:
        resmi_aylik_degisim = 0.0

    return {
        "df_analiz": df_analiz, "enf_genel": enf_genel, "enf_gida": enf_gida,
        "tahmin": month_end_forecast, "resmi_aylik_degisim": resmi_aylik_degisim,
        "son": son, "onceki_gun": onceki_gun, "gunler": gunler,
        "ad_col": ad_col, "agirlik_col": aktif_agirlik_col, "baz_col": baz_col, "gun_farki": gun_farki,
        "stats_urun": len(df_analiz), "stats_kategori": df_analiz['Grup'].nunique(),
        "stats_veri_noktasi": len(df_analiz) * len(tum_gunler_sirali)
    }

# 3. SIDEBAR UI
def ui_sidebar_ve_veri_hazirlama(df_analiz_base, raw_dates, ad_col):
    if df_analiz_base is None: return None
    st.sidebar.markdown("### ⚙️ Kontrol Paneli")
    
    # Lottie Animasyon (Varsa)
    lottie_url = "https://lottie.host/98606416-297c-4a37-9b2a-714013063529/5D6o8k8fW0.json"
    try:
        lottie_json = load_lottieurl(lottie_url)
        with st.sidebar:
             if lottie_json: st_lottie(lottie_json, height=100, key="nav_anim")
    except: pass

    BASLANGIC_LIMITI = "2026-02-04"
    tum_tarihler = sorted([d for d in raw_dates if d >= BASLANGIC_LIMITI], reverse=True)
    if not tum_tarihler:
        st.sidebar.warning("Veri henüz oluşmadı.")
        return None
    secilen_tarih = st.sidebar.selectbox("📅 Rapor Tarihi Seçin", options=tum_tarihler, index=0)
    
    st.sidebar.divider()
    
    st.sidebar.markdown("### 🌍 Piyasa Göstergeleri")
    symbols = [ 
        {"s": "FX_IDC:USDTRY", "d": "USD/TRY"}, 
        {"s": "FX_IDC:EURTRY", "d": "EUR/TRY"}, 
        {"s": "FX_IDC:XAUTRYG", "d": "Gram Altın (TL)"}, 
        {"s": "TVC:UKOIL", "d": "Brent Petrol"}, 
        {"s": "BINANCE:BTCUSDT", "d": "Bitcoin (USD)"} 
    ]
    # Widget'ları sadeleştir
    for sym in symbols:
        widget_code = f"""<div class="tradingview-widget-container" style="border-radius:8px; overflow:hidden; margin-bottom:8px; border:1px solid rgba(240, 246, 252, 0.1);"><div class="tradingview-widget-container__widget"></div><script type="text/javascript" src="https://s3.tradingview.com/external-embedding/embed-widget-mini-symbol-overview.js" async>{{ "symbol": "{sym['s']}", "width": "100%", "height": 70, "locale": "tr", "dateRange": "1D", "colorTheme": "dark", "isTransparent": true, "autosize": true, "largeChartUrl": "" }}</script></div>"""
        with st.sidebar: components.html(widget_code, height=72)
    
    # Tarih hesaplamaları
    tum_gunler_sirali = sorted([c for c in df_analiz_base.columns if re.match(r'\d{4}-\d{2}-\d{2}', str(c)) and c >= BASLANGIC_LIMITI])
    if secilen_tarih in tum_gunler_sirali:
        idx = tum_gunler_sirali.index(secilen_tarih)
        gunler = tum_gunler_sirali[:idx+1]
    else: gunler = tum_gunler_sirali
    if not gunler: return None
    son = gunler[-1]; dt_son = datetime.strptime(son, '%Y-%m-%d')
    col_w25, col_w26 = 'Agirlik_2025', 'Agirlik_2026'
    ZINCIR_TARIHI = datetime(2026, 2, 4)
    if dt_son >= ZINCIR_TARIHI:
        aktif_agirlik_col = col_w26
        gunler_2026 = [c for c in tum_gunler_sirali if c >= "2026-01-01"]
        baz_col = gunler_2026[0] if gunler_2026 else gunler[0]
    else:
        aktif_agirlik_col = col_w25; baz_col = gunler[0]

    ctx = hesapla_metrikler(df_analiz_base, secilen_tarih, gunler, tum_gunler_sirali, ad_col, agirlik_col=None, baz_col=baz_col, aktif_agirlik_col=aktif_agirlik_col, son=son)
    return ctx

# --- SAYFA FONKSİYONLARI (YENİ TASARIM) ---
def sayfa_ana_sayfa(ctx):
    # Modern, temiz hero section
    st.markdown(f"""
    <div style="text-align:left; padding: 40px 0; max-width: 900px;">
        <h1 style="font-size: 48px; font-weight: 800; letter-spacing: -1.5px; line-height: 1.1; margin-bottom: 24px;
                   background: linear-gradient(to right, #f0f6fc, #8b949e); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">
            Gerçek Zamanlı Enflasyon Analiz Terminali
        </h1>
        <p style="font-size: 18px; color: #8b949e; line-height: 1.6; margin-bottom: 40px;">
            Türkiye genelinde zincir marketler ve e-ticaret platformlarından toplanan yüksek frekanslı fiyat verileriyle oluşturulan alternatif, şeffaf ve yapay zeka destekli enflasyon göstergesi.
        </p>
        
        <div style="display:grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap:24px;">
            <div class="kpi-card">
                 <div style="display:flex; align-items:center; gap:10px; margin-bottom:10px;">
                    <span style="font-size:24px;">📦</span>
                    <span style="font-size:14px; font-weight:600; color:#8b949e;">TAKİP EDİLEN ÜRÜN</span>
                </div>
                <div class="kpi-value">{ctx["stats_urun"]}</div>
                <div style="font-size:13px; color:#2ea043;">Canlı veri akışı aktif</div>
            </div>
            <div class="kpi-card">
                 <div style="display:flex; align-items:center; gap:10px; margin-bottom:10px;">
                    <span style="font-size:24px;">📊</span>
                    <span style="font-size:14px; font-weight:600; color:#8b949e;">ALT SEKTÖR</span>
                </div>
                <div class="kpi-value">{ctx["stats_kategori"]}</div>
                 <div style="font-size:13px; color:#8b949e;">TÜİK ağırlıklandırması</div>
            </div>
            <div class="kpi-card">
                 <div style="display:flex; align-items:center; gap:10px; margin-bottom:10px;">
                    <span style="font-size:24px;">⚡</span>
                    <span style="font-size:14px; font-weight:600; color:#8b949e;">VERİ NOKTASI</span>
                </div>
                <div class="kpi-value">{ctx["stats_veri_noktasi"]}+</div>
                <div style="font-size:13px; color:#2563eb;">Yüksek frekanslı işlem</div>
            </div>
        </div>

        <div style="margin-top:40px; padding: 16px; background: rgba(37, 99, 235, 0.1); border: 1px solid rgba(37, 99, 235, 0.2); border-radius: 12px; display: inline-flex; align-items:center; gap:12px;">
            <span style="display:block; width:10px; height:10px; background:#2ea043; border-radius:50%;"></span>
            <span style="font-family: 'JetBrains Mono', monospace; font-size:13px; color: #f0f6fc;">
                SİSTEM DURUMU: NORMAL • SON GÜNCELLEME: {datetime.now().strftime('%H:%M')} TSİ
            </span>
        </div>
    </div>
    """, unsafe_allow_html=True)

def sayfa_piyasa_ozeti(ctx):
    st.markdown("### ⚡ Kritik Göstergeler")
    
    # KPI Kartları - Grid Yapısı
    c1, c2, c3, c4 = st.columns(4)
    
    def create_kpi(title, value, sub_icon, sub_text, trend_class):
        return f"""
        <div class="kpi-card">
            <div class="kpi-title">{title}</div>
            <div class="kpi-value">%{value:.2f}</div>
            <div class="kpi-sub {trend_class}">
                <span>{sub_icon}</span> <span>{sub_text}</span>
            </div>
        </div>
        """
        
    # Trend renklerini belirle
    enf_genel_trend = "trend-up" if ctx["enf_genel"] > 0 else "trend-down"
    enf_gida_trend = "trend-up" if ctx["enf_gida"] > 0 else "trend-down"

    with c1: st.markdown(create_kpi("GENEL ENFLASYON (AYLIK)", ctx["enf_genel"], "📈", "Sepet Değişimi", enf_genel_trend), unsafe_allow_html=True)
    with c2: st.markdown(create_kpi("GIDA ENFLASYONU (AYLIK)", ctx["enf_gida"], "🍲", "Mutfak Harcaması", enf_gida_trend), unsafe_allow_html=True)
    with c3: st.markdown(create_kpi("YIL SONU TAHMİNİ", ctx["tahmin"], "🤖", "AI Projeksiyonu", "trend-flat"), unsafe_allow_html=True)
    with c4: st.markdown(create_kpi("RESMİ TÜİK VERİSİ", ctx["resmi_aylik_degisim"], "🏛️", "Son Açıklanan", "trend-flat"), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    
    # Kayan Ticker Bandı (Daha sade)
    df = ctx["df_analiz"]
    inc = df.sort_values('Gunluk_Degisim', ascending=False).head(7)
    dec = df.sort_values('Gunluk_Degisim', ascending=True).head(7)
    ticker_items = []
    for _, r in inc.iterrows():
        if r['Gunluk_Degisim'] > 0: ticker_items.append(f"<span class='ticker-item trend-up'>▲ {r[ctx['ad_col']]} %{r['Gunluk_Degisim']*100:.1f}</span>")
    for _, r in dec.iterrows():
        if r['Gunluk_Degisim'] < 0: ticker_items.append(f"<span class='ticker-item trend-down'>▼ {r[ctx['ad_col']]} %{r['Gunluk_Degisim']*100:.1f}</span>")
        
    st.markdown(f"""
    <div class="ticker-container">
        <div class="ticker-track">
            {' • '.join(ticker_items)} {' • '.join(ticker_items)} </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Grafikler ve Özet Tablo
    col_g1, col_g2 = st.columns([2, 1])
    with col_g1:
        st.subheader("Fiyat Değişim Dağılımı")
        # Histogram için daha profesyonel renkler
        fig_hist = px.histogram(df, x="Fark_Yuzde", nbins=30, color_discrete_sequence=["#2563eb"])
        fig_hist.update_traces(marker_line_color='#161b22', marker_line_width=1, opacity=0.9)
        st.plotly_chart(style_chart(fig_hist), use_container_width=True)
        
    with col_g2:
        st.subheader("Piyasa Özeti")
        artan = len(df[df['Fark'] > 0])
        dusen = len(df[df['Fark'] < 0])
        sabit = len(df[df['Fark'] == 0])
        
        st.markdown(f"""
        <div style="background:var(--bg-surface-1); border-radius:16px; padding:24px; border:1px solid var(--border-subtle); height:100%; display:flex; flex-direction:column; justify-content:center;">
            <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:20px; padding-bottom:15px; border-bottom:1px solid var(--border-subtle);">
                <span style="font-size:14px; color:#8b949e; font-weight:600;">YÜKSELEN ÜRÜNLER</span>
                <span style="font-size:20px; color:#2ea043; font-weight:700; font-family:'JetBrains Mono';">▲ {artan}</span>
            </div>
            <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:20px; padding-bottom:15px; border-bottom:1px solid var(--border-subtle);">
                <span style="font-size:14px; color:#8b949e; font-weight:600;">DÜŞEN ÜRÜNLER</span>
                <span style="font-size:20px; color:#da3633; font-weight:700; font-family:'JetBrains Mono';">▼ {dusen}</span>
            </div>
             <div style="display:flex; justify-content:space-between; align-items:center;">
                <span style="font-size:14px; color:#8b949e; font-weight:600;">SABİT KALANLAR</span>
                <span style="font-size:20px; color:#f0f6fc; font-weight:700; font-family:'JetBrains Mono';">• {sabit}</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("Sektörel Isı Haritası (Ağırlıklı)")
    # Isı haritası için daha net bir renk skalası
    fig_tree = px.treemap(df, path=[px.Constant("Genel Pazar"), 'Grup', ctx['ad_col']], values=ctx['agirlik_col'], color='Fark', 
                          color_continuous_scale=['#da3633', '#161b22', '#2ea043'], # Kırmızı -> Koyu -> Yeşil
                          color_continuous_midpoint=0)
    fig_tree.update_traces(marker_line_width=1, marker_line_color='#161b22')
    st.plotly_chart(style_chart(fig_tree, is_sunburst=True), use_container_width=True)

def sayfa_kategori_detay(ctx):
    df = ctx["df_analiz"]
    st.markdown("### 🔍 Ürün Fiyat Takibi")
    
    # Filtreleme Alanı - Daha temiz layout
    with st.container():
        c1, c2 = st.columns([1, 3])
        kategoriler = ["Tümü"] + sorted(df['Grup'].unique().tolist())
        secilen_kat = c1.selectbox("Kategori Filtresi", kategoriler)
        arama = c2.text_input("Ürün Arama", placeholder="Örn: Süt, Ekmek, Yumurta...")
    
    st.divider()

    df_show = df.copy()
    if secilen_kat != "Tümü": df_show = df_show[df_show['Grup'] == secilen_kat]
    if arama: df_show = df_show[df_show[ctx['ad_col']].astype(str).str.contains(arama, case=False, na=False)]
    
    if not df_show.empty:
        # Pagination
        items_per_page = 20
        total_pages = max(1, len(df_show)//items_per_page + 1)
        
        col_pag, _ = st.columns([2, 8])
        page_num = col_pag.number_input(f"Sayfa (Toplam {total_pages})", min_value=1, max_value=total_pages, step=1)
        
        batch = df_show.iloc[(page_num - 1) * items_per_page : (page_num - 1) * items_per_page + items_per_page]
        
        # Grid Görünümü
        cols = st.columns(5) # 5 kolonlu grid
        for idx, row in enumerate(batch.to_dict('records')):
            fiyat = row[ctx['son']]; fark = row.get('Gunluk_Degisim', 0) * 100
            
            if fark > 0.01:
                badge_cls = "badge-up"; icon = "▲"; fark_txt = f"+%{abs(fark):.2f}"
            elif fark < -0.01:
                badge_cls = "badge-down"; icon = "▼"; fark_txt = f"-%{abs(fark):.2f}"
            else:
                badge_cls = "badge-flat"; icon = "•"; fark_txt = "%0.00"

            with cols[idx % 5]:
                st.markdown(f"""
                <div class="grid-card">
                    <div class="grid-title" title="{row[ctx['ad_col']]}">{row[ctx['ad_col']]}</div>
                    <div class="grid-price-row">
                        <div class="grid-price">{fiyat:.2f} ₺</div>
                        <div class="grid-badge {badge_cls}">{icon} {fark_txt}</div>
                    </div>
                </div>
                <div style="margin-bottom:16px;"></div>
                """, unsafe_allow_html=True)
    else: 
        st.info("Kriterlere uygun ürün bulunamadı.")

def sayfa_tam_liste(ctx):
    st.markdown("### 📋 Tam Veri Seti ve Analiz")
    df = ctx["df_analiz"]

    # Sparkline için veriyi hazırla
    def fix_sparkline(row):
        vals = [v if pd.notnull(v) else 0 for v in row.tolist()]
        if not vals: return [0,0]
        # Düz çizgi olmaması için minik bir varyasyon ekle (görsel hile)
        if min(vals) == max(vals) and len(vals) > 1: vals[-1] += 0.0001
        return vals

    df['Fiyat_Trendi'] = df[ctx['gunler']].apply(fix_sparkline, axis=1)
    
    cols_show = ['Grup', ctx['ad_col'], 'Fiyat_Trendi', ctx['baz_col'], ctx['son'], 'Gunluk_Degisim']
    
    cfg = {
        "Grup": st.column_config.TextColumn("Kategori", width="medium"),
        ctx['ad_col']: st.column_config.TextColumn("Ürün Adı", width="large"),
        "Fiyat_Trendi": st.column_config.LineChartColumn("Fiyat Trendi (Son Dönem)", width="medium", y_min=0), 
        ctx['baz_col']: st.column_config.NumberColumn(f"Baz Fiyat ({ctx['baz_col']})", format="%.2f ₺"), 
        ctx['son']: st.column_config.NumberColumn(f"Son Fiyat ({ctx['son']})", format="%.2f ₺"),
        "Gunluk_Degisim": st.column_config.ProgressColumn("Günlük Değişim", format="%.2f%%", min_value=-0.2, max_value=0.2), 
    }
    
    # Dataframe'i daha temiz göster
    st.dataframe(df[cols_show], column_config=cfg, hide_index=True, use_container_width=True, height=700)
    
    # İndirme butonu
    output = BytesIO(); 
    with pd.ExcelWriter(output) as writer: df.to_excel(writer, index=False)
    st.download_button("📥 Excel Raporunu İndir", data=output.getvalue(), file_name=f"Enflasyon_Verisi_{ctx['son']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

def sayfa_raporlama(ctx):
    st.markdown("### 📝 Stratejik Pazar Raporu (AI)")
    
    col_l, col_r = st.columns([7, 3])
    with col_l:
        # Rapor metnini bir "kağıt" gibi göster
        rap_text = generate_detailed_static_report(ctx["df_analiz"], ctx["son"], ctx["enf_genel"], ctx["enf_gida"], ctx["gun_farki"], ctx["tahmin"], ctx["ad_col"], ctx["agirlik_col"])
        st.markdown(f"""
        <div style="background:var(--bg-surface-1); padding:40px; border-radius:12px; border:1px solid var(--border-subtle); font-family:'Inter'; line-height:1.7; font-size:15px; color: var(--text-primary);">
            {rap_text.replace(chr(10), '<br>').replace('**', '<b style="color:var(--accent-blue);">').replace('**', '</b>')}
        </div>
        """, unsafe_allow_html=True)
        
    with col_r:
        st.markdown("#### Rapor İşlemleri")
        st.write("Bu rapor, sistemdeki anlık veriler ve yapay zeka algoritmaları kullanılarak otomatik oluşturulmuştur. Word formatında indirerek düzenleyebilirsiniz.")
        
        word_buffer = create_word_report(rap_text, ctx["son"], ctx["df_analiz"])
        st.download_button(
            label="📄 Word Raporu Olarak İndir", 
            data=word_buffer, 
            file_name=f"Piyasa_Raporu_{ctx['son']}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            use_container_width=True
        )
        
        st.divider()
        st.info("NOT: Bu belge, dahili analiz amaçlıdır ve resmi yatırım tavsiyesi niteliği taşımaz.")

def sayfa_maddeler(ctx):
    df = ctx["df_analiz"]
    st.markdown("### 📦 Kategori Bazlı Kümülatif Değişim")
    st.markdown("<p style='color:#8b949e; font-size:14px; margin-bottom:20px;'>Seçilen kategorideki ürünlerin, baz alınan tarihe göre toplam yüzdesel değişimlerini gösterir.</p>", unsafe_allow_html=True)
    
    kategoriler = sorted(df['Grup'].unique().tolist())
    col1, col2 = st.columns([1, 3])
    with col1: secilen_kat = st.selectbox("Kategori Seçiniz:", options=kategoriler, index=0)
    
    df_sub = df[df['Grup'] == secilen_kat].copy().sort_values('Fark_Yuzde', ascending=True)
    
    if not df_sub.empty:
        # Renkleri profesyonel palete göre ayarla
        colors = ['#da3633' if x < 0 else '#2ea043' for x in df_sub['Fark_Yuzde']]
        
        fig = go.Figure(go.Bar(
            x=df_sub['Fark_Yuzde'], 
            y=df_sub[ctx['ad_col']], 
            orientation='h', 
            marker_color=colors,
            text=df_sub['Fark_Yuzde'].apply(lambda x: f"%{x:.2f}"), 
            textposition='outside',
            hovertemplate='<b>%{y}</b><br>Kümülatif Değişim: %%{x:.2f}<extra></extra>'
        ))
        
        # Dinamik yükseklik
        fig.update_layout(
            height=max(600, len(df_sub) * 40), 
            title=f"{secilen_kat} Grubu Ürünleri",
            xaxis_title="Değişim Oranı (%)", 
            yaxis_title=None,
            margin=dict(l=0, r=0, t=40, b=0)
        )
        st.plotly_chart(style_chart(fig), use_container_width=True)
    else: 
        st.warning("Bu kategoride görüntülenecek veri bulunamadı.")

def sayfa_trend_analizi(ctx):
    st.markdown("### 📈 Zaman Serisi ve Trend Analizi")
    
    df = ctx["df_analiz"]; gunler = ctx["gunler"]; agirlik_col = ctx["agirlik_col"]
    
    # --- Genel Enflasyon Trendi ---
    endeks_verisi = []
    for gun in gunler:
        temp_df = df.dropna(subset=[gun, agirlik_col])
        if not temp_df.empty and temp_df[agirlik_col].sum() > 0:
            index_val = (temp_df[gun] * temp_df[agirlik_col]).sum() / temp_df[agirlik_col].sum()
            endeks_verisi.append({"Tarih": gun, "Deger": index_val})
            
    df_endeks = pd.DataFrame(endeks_verisi)
    if not df_endeks.empty:
        # Başlangıcı 0'a endeksle
        df_endeks['Kümülatif_Degisim'] = ((df_endeks['Deger'] / df_endeks.iloc[0]['Deger']) - 1) * 100
        
        fig_genel = px.line(df_endeks, x='Tarih', y='Kümülatif_Degisim', title="Genel Sepet Kümülatif Değişimi (%)", markers=True)
        # Çizgiyi kalınlaştır ve kurumsal mavi yap
        fig_genel.update_traces(line_color='#2563eb', line_width=4, marker_size=10)
        
        st.plotly_chart(style_chart(fig_genel), use_container_width=True)
        st.caption(f"ℹ️ Grafik, {gunler[0]} tarihini baz (0) alarak hesaplanan ağırlıklı sepet değişimini gösterir.")
    
    st.divider()
    
    # --- Ürün Bazlı Kıyaslama ---
    st.subheader("Ürün Bazlı Fiyat Kıyaslama")
    
    # Varsayılan olarak en çok artan 3 ürünü seç
    default_selection = df.sort_values('Fark_Yuzde', ascending=False).head(3)[ctx['ad_col']].tolist()
    seçilen_urunler = st.multiselect("Karşılaştırılacak ürünleri seçin:", options=df[ctx['ad_col']].unique(), default=default_selection)
    
    if seçilen_urunler:
        df_melted = df[df[ctx['ad_col']].isin(seçilen_urunler)][[ctx['ad_col']] + gunler].melt(id_vars=[ctx['ad_col']], var_name='Tarih', value_name='Fiyat')
        
        # Her ürünün kendi ilk günkü fiyatına göre değişimini hesapla
        base_prices = df_melted[df_melted['Tarih'] == gunler[0]].set_index(ctx['ad_col'])['Fiyat'].to_dict()
        df_melted['Yuzde_Degisim'] = df_melted.apply(lambda row: ((row['Fiyat']/base_prices.get(row[ctx['ad_col']], 1)) - 1)*100 if base_prices.get(row[ctx['ad_col']], 0) > 0 else 0, axis=1)
        
        fig_urun = px.line(df_melted, x='Tarih', y='Yuzde_Degisim', color=ctx['ad_col'], title="Seçili Ürünlerin Kümülatif Değişimi (%)", markers=True, color_discrete_sequence=px.colors.qualitative.Bold)
        fig_urun.update_traces(line_width=3)
        st.plotly_chart(style_chart(fig_urun), use_container_width=True)

def sayfa_metodoloji(ctx=None):
    st.markdown("### ℹ️ Metodoloji ve Yasal Uyarı")
    
    with st.expander("📌 Veri Toplama ve İşleme Yöntemi", expanded=True):
        st.markdown("""
        Bu sistem, Türkiye'de faaliyet gösteren önde gelen zincir marketler ve e-ticaret platformlarından **web kazıma (web scraping)** yöntemiyle günlük olarak fiyat verisi toplamaktadır.
        * **Kapsam:** TÜİK enflasyon sepetindeki ana harcama gruplarını temsil eden simüle edilmiş bir ürün sepeti takip edilmektedir.
        * **Veri Güvenliği:** Toplama işlemi sırasında User-Agent rotasyonu ve hız sınırlaması (rate limiting) uygulanarak hedef sitelere zarar verilmemesi amaçlanmaktadır.
        * **Veri Temizliği:** Aykırı değerler (anomali) ve eksik veriler, istatistiksel yöntemlerle (örneğin, önceki günün verisini taşıma veya enterpolasyon) işlenmektedir.
        """)

    with st.expander("🧮 Endeks Hesaplama Modeli", expanded=True):
        st.markdown("""
        Enflasyon hesaplamasında, uluslararası standartlara uygun olarak **Zincirleme Laspeyres Fiyat Endeksi** yaklaşımı benimsenmiştir.
        * **Formül:** `I(t) = Σ ( P(i,t) / P(i,0) ) × W(i)`
        * **Ağırlıklandırma (W):** Ürünlerin endeks üzerindeki etkisi, TÜİK Hanehalkı Bütçe Anketi'nden (HBA) elde edilen harcama paylarına göre simüle edilmiş ağırlıklarla belirlenir. Bu ağırlıklar her yılın başında (Zincirleme tarihi) güncellenir.
        """)

    with st.expander("⚠️ Yasal Uyarı ve Sorumluluk Reddi", expanded=True):
        st.markdown("""
        * Bu platformda sunulan veriler, analizler ve raporlar tamamen **bilgilendirme ve akademik/deneysel amaçlıdır.**
        * Buradaki bilgiler, **resmi enflasyon verisi (TÜİK) yerine geçmez** ve kesinlikle **yatırım tavsiyesi değildir.**
        * Verilerin doğruluğu, tamlığı veya güncelliği konusunda garanti verilmez. Bu verilere dayanarak alınan finansal veya ticari kararlardan doğabilecek zararlardan sistem geliştiricileri sorumlu tutulamaz.
        """)

# --- ANA UYGULAMA MANTIĞI ---
def main():
    # --- HEADER BÖLÜMÜ (Üstte sabit) ---
    c_head_1, c_head_2 = st.columns([3, 1])
    with c_head_1:
        st.markdown("""
        <div style="display:flex; align-items:center; gap:15px;">
            <div style="font-size:32px;">💠</div>
            <div>
                <div style="font-weight:800; font-size:26px; color:var(--text-primary); letter-spacing:-0.5px;">Piyasa Monitörü <span style="color:var(--accent-blue);">PRO</span></div>
                <div style="font-size:13px; color:var(--text-secondary);">Kurumsal Enflasyon Analiz Terminali v2.1</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with c_head_2:
        st.markdown(f"""
        <div style="text-align:right;">
            <div style="font-size:12px; font-weight:600; color:var(--text-secondary); letter-spacing:1px;">İSTANBUL (TSİ)</div>
            <div style="font-family:'JetBrains Mono'; font-size:22px; font-weight:700; color:var(--text-primary);">{datetime.now().strftime("%d.%m.%Y")}</div>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("<div style='margin-bottom:25px;'></div>", unsafe_allow_html=True)

    # --- NAVİGASYON ve AKSİYON BAR ---
    c_nav, c_act = st.columns([6, 1], gap="medium")
    
    with c_nav:
        # Modern Segmented Control Menü
        menu_items = {
            "🏠 Ana Sayfa": "Ana Sayfa",
            "⚡ Özet": "Piyasa Özeti",
            "📈 Trendler": "Trendler",
            "📦 Ürünler": "Maddeler",
            "🔍 Detay": "Kategori Detay",
            "💾 Veri Seti": "Tam Liste",
            "📝 Rapor": "Raporlama",
            "ℹ️ Metodoloji": "Metodoloji"
        }
        # CSS ile özelleştirilmiş radio butonu
        secilen_etiket = st.radio(
            "Navigasyon", 
            options=list(menu_items.keys()), 
            label_visibility="collapsed", 
            key="main_nav",
            horizontal=True
        )
        secim = menu_items[secilen_etiket]

    with c_act:
        # Senkronizasyon Butonu
        if st.button("🔄 Veriyi Güncelle", use_container_width=True, help="Github üzerinden son verileri çeker ve işler."):
            with st.status("Veri senkronizasyonu başlatılıyor...", expanded=True) as status:
                st.write("Bot sunucusuna bağlanılıyor...")
                progress_bar = st.progress(0)
                res = html_isleyici(lambda p: progress_bar.progress(min(1.0, max(0.0, p))))
                
                if "OK" in res:
                    progress_bar.progress(1.0)
                    status.update(label="Senkronizasyon tamamlandı!", state="complete", expanded=False)
                    st.cache_data.clear()
                    st.toast('Sistem başarıyla güncellendi.', icon='✅')
                    time.sleep(1)
                    st.rerun()
                elif "Veri bulunamadı" in res:
                    status.update(label="Yeni veri yok.", state="complete")
                    st.toast("İşlenecek yeni veri paketi bulunamadı.", icon="ℹ️")
                else:
                    status.update(label="Hata oluştu!", state="error")
                    st.error(f"Senkronizasyon hatası: {res}")

    st.divider()

    # --- VERİ YÜKLEME ---
    with st.spinner("Veritabanı ve analiz motoru başlatılıyor..."):
        df_base, r_dates, col_name = verileri_getir_cache()
    
    if df_base is not None:
        # Sidebar'ı sadece veri varsa göster
        ctx = ui_sidebar_ve_veri_hazirlama(df_base, r_dates, col_name)
    else:
        ctx = None

    # --- SAYFA İÇERİĞİ ---
    if ctx:
        # İçerik alanı için biraz boşluk
        st.markdown("<div style='margin-top:20px;'></div>", unsafe_allow_html=True)
        
        if secim == "Ana Sayfa": sayfa_ana_sayfa(ctx)
        elif secim == "Piyasa Özeti": sayfa_piyasa_ozeti(ctx)
        elif secim == "Trendler": sayfa_trend_analizi(ctx)
        elif secim == "Maddeler": sayfa_maddeler(ctx)
        elif secim == "Kategori Detay": sayfa_kategori_detay(ctx)
        elif secim == "Tam Liste": sayfa_tam_liste(ctx)
        elif secim == "Raporlama": sayfa_raporlama(ctx)
        elif secim == "Metodoloji": sayfa_metodoloji(ctx)
        
    else:
        # Veri yüklenemezse
        if secim == "Metodoloji":
            sayfa_metodoloji()
        else:
            st.error("⚠️ Veritabanı bağlantısı kurulamadı. Lütfen internet bağlantınızı kontrol edin veya GitHub yapılandırmasını gözden geçirin.")

    # --- FOOTER ---
    st.markdown("""
    <div style="text-align:center; color:var(--text-secondary); font-size:12px; margin-top:60px; padding-top:20px; border-top:1px solid var(--border-subtle);">
        VALIDASYON MÜDÜRLÜĞÜ © 2026 • CONFIDENTIAL ANALYTICS SUITE • v2.1.4 (Stable)
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
